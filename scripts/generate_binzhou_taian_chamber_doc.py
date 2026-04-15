from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE_DIR = Path(r"C:\Users\weixi\Desktop\5564~1")
OUTPUT_PATH = BASE_DIR / "滨州泰安商会揭牌仪式策划方案.docx"

TITLE_FONT = "方正小标宋简体"
BODY_FONT = "仿宋"
HEADING_FONT = "黑体"
EMPHASIS_FONT = "楷体_GB2312"


def set_run_font(run, font_name: str, size_pt: float):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)


def configure_paragraph(paragraph, *, first_line_indent_pt: float | None = None, align=None):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(15.5)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if first_line_indent_pt is not None:
        fmt.first_line_indent = Pt(first_line_indent_pt)
    if align is not None:
        paragraph.alignment = align


def add_title_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    configure_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    set_run_font(run, TITLE_FONT, 16)


def add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    configure_paragraph(p, first_line_indent_pt=30)
    run = p.add_run(text)
    set_run_font(run, BODY_FONT, 15)


def add_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    configure_paragraph(p)
    run = p.add_run(text)
    set_run_font(run, HEADING_FONT, 15)


def add_label_paragraph(doc: Document, label: str, content: str):
    p = doc.add_paragraph()
    configure_paragraph(p)
    r1 = p.add_run(label)
    set_run_font(r1, EMPHASIS_FONT, 15)
    r2 = p.add_run(content)
    set_run_font(r2, BODY_FONT, 15)


def add_subitem_title(doc: Document, text: str):
    p = doc.add_paragraph()
    configure_paragraph(p)
    run = p.add_run(text)
    set_run_font(run, EMPHASIS_FONT, 15)


def set_cell_paragraph(cell, text: str, *, bold: bool = False, font_name: str = BODY_FONT, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    configure_paragraph(p, align=align)
    run = p.add_run(text)
    run.bold = bold
    set_run_font(run, font_name, 12)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_document_layout(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    style.font.size = Pt(15)

    section = doc.sections[0]
    section.top_margin = Cm(2.8)
    section.bottom_margin = Cm(2.7)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.7)


def add_video_project_table(doc: Document):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows = [
        ("项目", "内容"),
        ("标题", "滨州泰安商会揭牌仪式"),
        ("时长", "活动全流程记录与一分钟以内快剪短视频"),
        ("内容", "现场揭牌环节、嘉宾签到、主视觉展示、主持串场、氛围节目及合影留念等"),
        ("比例", "16:9"),
        ("画质", "1080P"),
        ("配乐", "根据现场氛围统一搭配"),
        ("格式", "MP4"),
        ("预计交付时间", "2026年10月18日至2026年10月20日"),
    ]
    for row_idx, (a, b) in enumerate(rows):
        cells = table.add_row().cells
        set_cell_paragraph(cells[0], a, bold=row_idx == 0, font_name=HEADING_FONT if row_idx == 0 else BODY_FONT)
        set_cell_paragraph(cells[1], b, bold=row_idx == 0, font_name=HEADING_FONT if row_idx == 0 else BODY_FONT)


def add_material_table(doc: Document):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows = [
        ("类别", "数量"),
        ("签名墙", "1套"),
        ("道旗", "10个"),
        ("空飘", "8个"),
        ("竖幅", "38个"),
        ("横幅", "2条"),
        ("合影台", "1套"),
        ("摄影师", "1人"),
        ("讲话台鲜花", "1套"),
        ("舞台地毯", "100平方"),
        ("户外花篮", "12个"),
        ("水鼓舞", "6人"),
        ("揭牌用红绸子", "2套"),
        ("定制礼品", "100套"),
    ]
    for row_idx, (a, b) in enumerate(rows):
        cells = table.add_row().cells
        set_cell_paragraph(cells[0], a, bold=row_idx == 0, font_name=HEADING_FONT if row_idx == 0 else BODY_FONT)
        set_cell_paragraph(cells[1], b, bold=row_idx == 0, font_name=HEADING_FONT if row_idx == 0 else BODY_FONT)


def add_flow_table(doc: Document):
    table = doc.add_table(rows=0, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows = [
        ("时间", "内容", "视听"),
        ("17:00-17:40", "嘉宾签到、签到墙合影、领取纪念品、引导入场就座", "屏幕播放活动主视觉，配合舒缓背景音乐"),
        ("17:40-18:00", "暖场播放、嘉宾就座、主持准备开场", "循环播放预热短视频与现场暖场素材"),
    ]
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 or col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_paragraph(
                cells[col_idx],
                value,
                bold=row_idx == 0,
                font_name=HEADING_FONT if row_idx == 0 else BODY_FONT,
                align=align,
            )


def add_budget_table(doc: Document):
    table = doc.add_table(rows=0, cols=8)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows = [
        ("序号", "项目", "规格", "数量", "单位", "单价", "合计", "备注"),
        ("1", "签名墙", "桁架3m×5m", "1", "套", "600", "600", ""),
        ("2", "签名墙", "喷绘3m×7m", "21", "平方", "38", "798", ""),
        ("3", "道旗", "", "10", "个", "150", "1500", "罗马高档底座"),
        ("4", "空飘", "", "8", "个", "200", "1600", "1个气球带1条竖幅"),
        ("5", "竖幅", "0.9×26m", "38", "个", "150", "5700", ""),
        ("6", "横幅", "", "2", "条", "60", "120", "1号厅、3号厅"),
        ("7", "合影台", "", "1", "套", "700", "700", "满足拍照需要"),
        ("8", "摄影师", "", "1", "人", "800", "800", ""),
        ("9", "讲话台鲜花", "", "1", "套", "200", "200", ""),
        ("10", "舞台地毯", "20×5", "100", "平方", "7", "700", ""),
        ("11", "户外花篮", "", "12", "个", "100", "1200", ""),
        ("12", "水鼓舞", "", "6", "人", "300", "1800", ""),
        ("13", "拱门", "", "1", "个", "260", "260", "宽15m、带横幅"),
        ("14", "揭牌用红绸子", "", "2", "套", "60", "120", ""),
        ("15", "主持人", "", "1", "人", "1500", "1500", "成立和晚会"),
        ("16", "签名笔", "", "10", "", "", "0", "赞助"),
        ("17", "如意标签", "", "2", "", "", "0", "赞助"),
        ("18", "地贴", "", "16", "", "", "0", "赞助"),
        ("19", "嘉宾证", "", "", "", "", "0", "赞助"),
        ("20", "讲话台", "", "1", "", "", "0", "赞助"),
        ("", "合计", "", "", "", "", "17598", ""),
        ("", "实际收取", "", "", "", "", "17500", ""),
    ]
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(values):
            set_cell_paragraph(
                cells[col_idx],
                value,
                bold=row_idx == 0,
                font_name=HEADING_FONT if row_idx == 0 else BODY_FONT,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )


def add_right_signature(doc: Document, text: str):
    p = doc.add_paragraph()
    configure_paragraph(p, align=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(text)
    set_run_font(run, BODY_FONT, 15)


def build_document():
    doc = Document()
    set_document_layout(doc)

    add_title_paragraph(doc, "滨州泰安商会")
    add_title_paragraph(doc, "揭牌仪式策划方案")

    doc.add_paragraph()
    add_body_paragraph(
        doc,
        "为做好滨州泰安商会揭牌仪式的筹备与执行工作，统筹现场搭建、流程安排、影像记录和预算控制，"
        "现结合活动实际需要，形成本策划方案，以便统一执行标准，明确工作内容，提高现场落地效率。"
    )
    add_body_paragraph(
        doc,
        "本方案围绕活动基础信息、执行内容、物料配置、流程安排和预算明细等方面展开，"
        "用于揭牌仪式前期筹备、现场组织及后续对账。"
    )

    doc.add_paragraph()
    add_heading(doc, "一、活动概述")
    add_label_paragraph(doc, "活动名称：", "滨州泰安商会揭牌仪式。")
    add_label_paragraph(doc, "活动定位：", "以揭牌仪式为核心，兼顾商务接待、现场氛围营造、影像记录及传播展示。")
    add_label_paragraph(doc, "活动目标：", "保障揭牌仪式庄重有序，现场流程流畅，物料配置完整，费用核算清晰。")

    doc.add_paragraph()
    add_heading(doc, "二、基础信息")
    add_label_paragraph(doc, "活动形式：", "揭牌仪式、嘉宾签到、合影留念、节目暖场及现场主持串联。")
    add_label_paragraph(doc, "活动内容：", "包含签名墙布置、道旗与横竖幅展示、合影台搭建、舞台地毯铺设、主持与摄影执行等。")
    add_label_paragraph(doc, "现场重点：", "突出揭牌环节的仪式感，兼顾嘉宾动线、合影区域、主视觉展示和现场氛围。")

    doc.add_paragraph()
    add_heading(doc, "三、执行内容")
    add_subitem_title(doc, "1、现场影像执行")
    add_body_paragraph(
        doc,
        "活动现场采用固定机位进行全景记录，并同步完成重点环节拍摄，确保揭牌、合影和主持串场等核心内容完整留存。"
    )
    add_body_paragraph(
        doc,
        "活动结束后输出现场全程记录视频及一分钟以内快剪短视频，满足资料存档与后续传播使用。"
    )

    add_subitem_title(doc, "2、现场氛围执行")
    add_body_paragraph(
        doc,
        "通过签名墙、道旗、横幅、竖幅、合影台、讲话台鲜花、舞台地毯等内容，构建统一的活动视觉氛围。"
    )
    add_body_paragraph(
        doc,
        "围绕揭牌环节配置红绸道具与主持串联，增强现场仪式感和拍摄效果。"
    )

    add_subitem_title(doc, "3、影像交付安排")
    add_video_project_table(doc)

    doc.add_paragraph()
    add_heading(doc, "四、活动物料")
    add_material_table(doc)

    doc.add_paragraph()
    add_heading(doc, "五、活动流程")
    add_flow_table(doc)

    doc.add_paragraph()
    add_heading(doc, "六、活动预算")
    add_budget_table(doc)

    doc.add_paragraph()
    add_right_signature(doc, "滨州泰安商会揭牌仪式筹备组")
    add_right_signature(doc, "2026年4月10日")

    doc.save(str(OUTPUT_PATH))
    print("generated")


if __name__ == "__main__":
    build_document()
